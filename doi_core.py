"""
Crossref DOI verification for a reference list. No UI - import from app.py,
a notebook, or a script.

Two checks per reference:
  1. BACKWARD - resolve the DOI already in the entry and see which paper it
     actually points to. Catches a real DOI registered to a different article.
  2. FORWARD  - bibliographic search on the whole reference string, taking the
     best-matching record as the suggested correct DOI.

Callers set EMAIL and SCORE_THRESHOLD as module attributes before calling.
"""

import io
import re
import time
import difflib
import requests
import pandas as pd

try:
    import docx
except ImportError:
    import subprocess
    import sys
    subprocess.run([sys.executable, "-m", "pip", "install", "-q", "python-docx"])
    import docx

from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor


# ----------------------------------------------------------------------
# 1. INPUT
# ----------------------------------------------------------------------

# Path to your Word file. Leave "" to be prompted for an upload in Colab.
# Set USE_DOCX = False to check the pasted list below instead.
USE_DOCX = True
DOCX_PATH = ""

# Used only when USE_DOCX = False. One reference per line.
REFERENCES_TEXT = """
Sung, H., Ferlay, J., Siegel, R. L. (2021). Global cancer statistics 2020: GLOBOCAN estimates of incidence and mortality worldwide for 36 cancers in 185 countries. CA: A Cancer Journal for Clinicians, 71(3), 209-249. https://doi.org/10.3322/caac.21660
"""

# OPTIONAL. Crossref's API is free and needs no key. Supplying an email puts
# your requests in their "polite pool" - better provisioned, less likely to be
# throttled on long lists. Leave it as "" to skip it entirely; the only cost is
# the shared anonymous pool, which can be slower. Nothing is sent anywhere else.
EMAIL = ""

# Below this score a suggested DOI is reported but flagged as unconfirmed
SCORE_THRESHOLD = 0.65

# Seconds between API calls - be gentle, Crossref is free
DELAY = 0.4

CSV_OUT = "doi_check_results.csv"

# Highlight entries that pass too (green). Set False to mark only the problems.
HIGHLIGHT_OK = True

# Leave runs that already carry a highlight untouched. Copyeditors and reviewers
# often highlight in yellow; overwriting it destroys their marks.
PRESERVE_EXISTING_HIGHLIGHT = True

# Headings that open a reference section (add your own if needed)
REF_HEADINGS = [
    "references", "reference", "bibliography", "works cited",
    "daftar pustaka", "literature cited", "reference list",
]

# Headings that close it
STOP_HEADINGS = [
    "appendix", "appendices", "acknowledgement", "acknowledgements",
    "acknowledgment", "acknowledgments", "supplementary material",
    "supplementary", "author contributions", "funding",
]


# ----------------------------------------------------------------------
# 2. TEXT HELPERS
# ----------------------------------------------------------------------

STOPWORDS = {
    "a", "an", "and", "the", "of", "in", "on", "for", "to", "with", "from",
    "by", "at", "as", "is", "are", "be", "using", "via", "into", "its",
}

DOI_PATTERN = re.compile(r"10\.\d{4,9}/[^\s\"<>,;]+", re.IGNORECASE)


def normalize(text):
    """Lowercase, punctuation to spaces, collapse whitespace."""
    lowered = text.lower()
    cleaned = ""
    for char in lowered:
        if char.isalnum() or char.isspace():
            cleaned += char
        else:
            cleaned += " "
    return " ".join(cleaned.split())


def tokenize(text):
    """Content words only, duplicates removed."""
    tokens = []
    for word in normalize(text).split():
        if len(word) > 2 and word not in STOPWORDS and word not in tokens:
            tokens.append(word)
    return tokens


def extract_doi(reference):
    """Pull a DOI out of a reference string, or return '' if there is none."""
    match = DOI_PATTERN.search(reference)
    if match is None:
        return ""
    doi = match.group(0)
    while len(doi) > 0 and doi[-1] in ".,;)]}":
        doi = doi[:-1]
    return doi.lower()


def extract_year(reference):
    """First 19xx/20xx found in the reference, as a string."""
    match = re.search(r"(19|20)\d{2}", reference)
    if match is None:
        return ""
    return match.group(0)


# ----------------------------------------------------------------------
# 3. SCORING - does this Crossref record match this reference string?
# ----------------------------------------------------------------------

def surname_is_close(surname, ref_start):
    """
    Catch corrupted or truncated surnames - a reference manager that imported
    'Garcia-Lara' as 'Garci' should still count as a partial author match.
    """
    for part in surname.split():
        if len(part) < 4:
            continue
        for word in ref_start.split():
            if len(word) < 4:
                continue
            if difflib.SequenceMatcher(None, part, word).ratio() >= 0.7:
                return True
    return False


def score_match(reference, record):
    """
    Return (score 0-1, note). Score combines title-word coverage, lead author
    surname, and year. A year conflict caps the score - it is the strongest
    single sign that a DOI belongs to a different article.
    """
    ref_norm = normalize(reference)
    notes = []

    title = record.get("title", "")
    title_tokens = tokenize(title)
    hits = 0
    for token in title_tokens:
        if token in ref_norm:
            hits += 1
    if len(title_tokens) == 0:
        title_score = 0.0
    else:
        title_score = hits / len(title_tokens)

    # fallback for short or reworded titles
    fuzzy = difflib.SequenceMatcher(None, normalize(title), ref_norm).ratio()
    title_score = max(title_score, fuzzy * 0.8)

    # the lead author should sit near the front of an APA-style reference;
    # finding the surname only deep in the string is weaker evidence
    surname = normalize(record.get("first_author", ""))
    author_missing = False
    if surname == "":
        author_score = 0.5
        notes.append("no author in record")
    elif surname in ref_norm[:60]:
        author_score = 1.0
    elif surname in ref_norm:
        author_score = 0.6
        notes.append("lead author is not first in your entry")
    elif surname_is_close(surname, ref_norm[:80]):
        author_score = 0.8
        notes.append("lead author spelt differently in your entry - check it")
    else:
        author_score = 0.0
        author_missing = True
        notes.append("first author differs")

    ref_year = extract_year(reference)
    rec_year = record.get("year", "")
    year_conflict = False
    if ref_year == "" or rec_year == "":
        year_score = 0.5
    elif abs(int(ref_year) - int(rec_year)) <= 1:
        year_score = 1.0
    else:
        year_score = 0.0
        year_conflict = True
        notes.append("year " + rec_year + " vs " + ref_year)

    score = 0.65 * title_score + 0.20 * author_score + 0.15 * year_score

    if year_conflict:
        score = min(score, 0.40)

    # A near-perfect title match identifies the paper on its own. Don't let a
    # corrupted author field in the user's own entry veto it - flag it instead.
    if author_missing and title_score < 0.90:
        score = min(score, 0.55)

    return round(score, 3), "; ".join(notes)


# ----------------------------------------------------------------------
# 4. CROSSREF CALLS
# ----------------------------------------------------------------------

def parse_record(item):
    """Flatten a Crossref `message` item into the few fields we need."""
    title_list = item.get("title", [])
    if len(title_list) > 0:
        title = title_list[0]
    else:
        title = ""

    container = item.get("container-title", [])
    if len(container) > 0:
        journal = container[0]
    else:
        journal = ""

    first_author = ""
    authors = item.get("author", [])
    if len(authors) > 0:
        first_author = authors[0].get("family", "")

    year = ""
    issued = item.get("issued", {}).get("date-parts", [[]])
    if len(issued) > 0 and len(issued[0]) > 0 and issued[0][0] is not None:
        year = str(issued[0][0])

    return {
        "doi": str(item.get("DOI", "")).lower(),
        "title": title,
        "journal": journal,
        "first_author": first_author,
        "year": year,
    }


def get_json(url, params):
    """GET with one retry. Returns the parsed JSON dict, or None."""
    params = dict(params)
    if EMAIL.strip() == "":
        params.pop("mailto", None)
        headers = {"User-Agent": "doi-check/1.0"}
    else:
        headers = {"User-Agent": "doi-check/1.0 (mailto:" + EMAIL + ")"}
    for attempt in range(2):
        try:
            response = requests.get(url, params=params, headers=headers, timeout=25)
        except Exception as error:
            print("   request failed:", error)
            time.sleep(2)
            continue
        if response.status_code == 200:
            return response.json()
        if response.status_code == 404:
            return None
        time.sleep(2)
    return None


def resolve_doi(doi):
    """Backward check: what does this DOI actually point to? None if unregistered."""
    url = "https://api.crossref.org/works/" + doi
    payload = get_json(url, {"mailto": EMAIL})
    if payload is None:
        return None
    return parse_record(payload["message"])


def search_reference(reference, rows=3):
    """Forward check: best Crossref candidates for this reference string."""
    url = "https://api.crossref.org/works"
    params = {
        "query.bibliographic": reference[:400],
        "rows": rows,
        "select": "DOI,title,container-title,author,issued",
        "mailto": EMAIL,
    }
    payload = get_json(url, params)
    if payload is None:
        return []
    candidates = []
    for item in payload["message"]["items"]:
        candidates.append(parse_record(item))
    return candidates


def best_candidate(reference, candidates):
    """Pick the highest-scoring candidate. Returns (record, score, note)."""
    best_record = None
    best_score = -1.0
    best_note = ""
    for record in candidates:
        score, note = score_match(reference, record)
        if score > best_score:
            best_record = record
            best_score = score
            best_note = note
    if best_record is None:
        return None, 0.0, ""
    return best_record, best_score, best_note


# ----------------------------------------------------------------------
# 5. MAIN CHECK
# ----------------------------------------------------------------------

def check_reference(reference):
    """Run both checks on one reference and return a result row."""
    given_doi = extract_doi(reference)

    row = {
        "reference": reference,
        "given_doi": given_doi,
        "status": "",
        "correct_doi": "",
        "resolves_to": "",
        "score": 0.0,
        "note": "",
    }

    # ---- backward check -------------------------------------------------
    if given_doi != "":
        record = resolve_doi(given_doi)
        time.sleep(DELAY)

        if record is None:
            row["status"] = "DOI DOES NOT EXIST"
            row["note"] = "not registered at Crossref - almost certainly fabricated"
        else:
            score, note = score_match(reference, record)
            row["resolves_to"] = record["first_author"] + " (" + record["year"] + ") " + record["title"]
            row["score"] = score
            if score >= SCORE_THRESHOLD:
                row["status"] = "OK"
                row["correct_doi"] = given_doi
                row["note"] = note
                return row
            row["status"] = "WRONG PAPER"
            row["note"] = "resolves to a different article; " + note
    else:
        row["status"] = "NO DOI"

    # ---- forward check --------------------------------------------------
    candidates = search_reference(reference)
    time.sleep(DELAY)
    record, score, note = best_candidate(reference, candidates)

    if record is None:
        row["note"] = (row["note"] + "; no Crossref candidates").strip("; ")
        return row

    row["correct_doi"] = record["doi"]
    row["score"] = score
    suggestion = record["first_author"] + " (" + record["year"] + ") " + record["title"]

    if score < SCORE_THRESHOLD:
        row["status"] = row["status"] + " - UNCONFIRMED"
        row["note"] = ("best guess only, verify by hand: " + suggestion + "; " + note).strip("; ")
    else:
        row["note"] = ("suggested: " + suggestion + "; " + note).strip("; ")

    return row


def check_all(references, on_progress=None):
    """
    references = list of [text, [paragraph indices]].
    on_progress(index, total, row) is called after each lookup so a UI can
    report progress; omit it and nothing is printed.
    """
    rows = []
    for i, entry in enumerate(references):
        row = check_reference(entry[0])
        row["n"] = i + 1
        row["paragraphs"] = entry[1]
        rows.append(row)
        if on_progress is not None:
            on_progress(i + 1, len(references), row)
    return pd.DataFrame(rows)


# ----------------------------------------------------------------------
# 6. READING THE DOCX
# ----------------------------------------------------------------------

def is_heading(paragraph):
    """Heading by style, or a short standalone line with no sentence punctuation."""
    style = ""
    if paragraph.style is not None and paragraph.style.name is not None:
        style = paragraph.style.name.lower()
    if style.startswith("heading") or style == "title":
        return True
    text = paragraph.text.strip()
    if len(text) == 0 or len(text) > 40:
        return False
    if text.endswith("."):
        return False
    return normalize(text) in REF_HEADINGS + STOP_HEADINGS


def strip_numbering(text):
    """Remove a leading '1.', '[12]', '(3)' so the query starts at the author."""
    return re.sub(r"^\s*[\[\(]?\d{1,3}[\]\)\.]\s*", "", text).strip()


def read_references_from_docx(path):
    """
    Return (list of (reference_text, [paragraph indices]), start_index).
    A paragraph with no year is treated as the wrapped tail of the one above it.
    """
    document = docx.Document(path)
    paragraphs = document.paragraphs

    start = -1
    for i, paragraph in enumerate(paragraphs):
        if normalize(paragraph.text) in REF_HEADINGS:
            start = i
            break
    if start < 0:
        raise ValueError(
            "No reference heading found. Expected one of: " + ", ".join(REF_HEADINGS)
            + ". Add yours to REF_HEADINGS."
        )

    references = []
    for i in range(start + 1, len(paragraphs)):
        paragraph = paragraphs[i]
        text = paragraph.text.strip()

        if text == "":
            continue
        if is_heading(paragraph) and normalize(text) in STOP_HEADINGS:
            break

        text = strip_numbering(text)

        is_continuation = extract_year(text) == "" and len(references) > 0 and len(text) < 200
        if is_continuation:
            references[-1][0] = references[-1][0] + " " + text
            references[-1][1].append(i)
        else:
            if len(text) < 25:
                continue
            references.append([text, [i]])

    return document, references


def load_references_from_text(text):
    """One reference per line; blank lines ignored. Same shape as the docx reader."""
    references = []
    for line in text.strip().split("\n"):
        stripped = line.strip()
        if stripped != "":
            references.append([strip_numbering(stripped), []])
    return references


# ----------------------------------------------------------------------
# 7. MARKING THE DOCX
# ----------------------------------------------------------------------

def status_colour(status):
    """(highlight, note colour, symbol) for a status string.

    Pink = broken, yellow = needs attention, green = verified. Word's RED
    highlight is too dark to read black text on, so pink carries the errors.
    """
    if status == "OK":
        return WD_COLOR_INDEX.BRIGHT_GREEN, RGBColor(0, 110, 0), "OK"
    if status.startswith("WRONG PAPER"):
        return WD_COLOR_INDEX.PINK, RGBColor(192, 0, 0), "WRONG PAPER"
    if status.startswith("DOI DOES NOT EXIST"):
        return WD_COLOR_INDEX.PINK, RGBColor(192, 0, 0), "DOI NOT REGISTERED"
    if status.startswith("NO DOI"):
        return WD_COLOR_INDEX.YELLOW, RGBColor(160, 100, 0), "NO DOI"
    return WD_COLOR_INDEX.YELLOW, RGBColor(160, 100, 0), "CHECK"


def build_note(row):
    """The short inline annotation appended after a reference."""
    _, _, symbol = status_colour(row["status"])
    parts = ["  [" + symbol]

    if row["status"] == "OK":
        parts.append("]")
        return "".join(parts)

    if row["resolves_to"] != "":
        parts.append(" - given DOI points to: " + row["resolves_to"][:90])
    if row["correct_doi"] != "":
        parts.append(" | use: " + row["correct_doi"])
    if row["status"].endswith("UNCONFIRMED"):
        parts.append(" (unconfirmed, verify by hand)")
    parts.append("]")
    return "".join(parts)


def mark_docx(document, results):
    """
    Highlight each reference paragraph by status, append an inline note, and
    return the finished .docx as bytes ready for a download button.
    """
    paragraphs = document.paragraphs

    for _, row in results.iterrows():
        indices = row["paragraphs"]
        if len(indices) == 0:
            continue

        highlight, colour, _ = status_colour(row["status"])

        if row["status"] != "OK" or HIGHLIGHT_OK:
            for index in indices:
                for run in paragraphs[index].runs:
                    if run.font.highlight_color is not None and PRESERVE_EXISTING_HIGHLIGHT:
                        continue
                    run.font.highlight_color = highlight

        note_run = paragraphs[indices[-1]].add_run(build_note(row))
        note_run.font.color.rgb = colour
        note_run.font.size = Pt(8)
        note_run.font.bold = True
        note_run.font.highlight_color = None

    try:
        append_summary(document, results)
    except Exception as error:
        print("   summary block skipped (" + str(error) + ") - highlighting is unaffected")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def has_style(document, name):
    """python-docx only sees styles present in the file's template."""
    try:
        document.styles[name]
        return True
    except KeyError:
        return False


def add_bullet(document, text, size=9):
    """Bullet that survives documents without a 'List Bullet' style."""
    if has_style(document, "List Bullet"):
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(text)
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run("\u2022  " + text)
    run.font.size = Pt(size)
    return paragraph


def append_summary(document, results):
    """A short summary block at the end of the document."""
    document.add_page_break()
    heading = document.add_paragraph()
    heading_run = heading.add_run("DOI check summary")
    heading_run.font.bold = True
    heading_run.font.size = Pt(14)

    counts = {}
    for status in results["status"]:
        key = status.split(" - ")[0]
        if key in counts:
            counts[key] = counts[key] + 1
        else:
            counts[key] = 1

    line = document.add_paragraph()
    summary_run = line.add_run(str(len(results)) + " references checked against Crossref")
    summary_run.font.size = Pt(9)

    for key in counts:
        add_bullet(document, key + ": " + str(counts[key]))

    for _, row in results.iterrows():
        if row["status"] == "OK":
            continue
        entry = document.add_paragraph()
        label = entry.add_run("[" + str(row["n"]) + "] " + row["status"] + " - ")
        label.font.bold = True
        label.font.size = Pt(9)
        label.font.color.rgb = RGBColor(192, 0, 0)
        body = entry.add_run(row["reference"][:120])
        body.font.size = Pt(9)
        if row["correct_doi"] != "":
            fix = entry.add_run("  ->  https://doi.org/" + row["correct_doi"])
            fix.font.size = Pt(9)
            fix.font.bold = True


# ----------------------------------------------------------------------
# 8. CONSOLE REPORT
# ----------------------------------------------------------------------

def report(results):
    """Print a compact summary of everything that needs attention."""
    print("\n" + "=" * 70)
    print("SUMMARY")
    print("=" * 70)
    for status in ["OK", "WRONG PAPER", "DOI DOES NOT EXIST", "NO DOI"]:
        count = 0
        for value in results["status"]:
            if value.startswith(status):
                count += 1
        if count > 0:
            print(str(count).rjust(3) + "  " + status)

    print("\nNEEDS FIXING")
    print("-" * 70)
    problems = 0
    for _, row in results.iterrows():
        if row["status"] == "OK":
            continue
        problems += 1
        print("\n[" + str(row["n"]) + "] " + row["status"])
        print("  reference : " + row["reference"][:100])
        if row["given_doi"] != "":
            print("  given     : " + row["given_doi"])
        if row["resolves_to"] != "":
            print("  points to : " + row["resolves_to"][:100])
        if row["correct_doi"] != "":
            print("  use       : " + row["correct_doi"] + "   (score " + str(row["score"]) + ")")
            print("  check     : https://doi.org/" + row["correct_doi"])
        if row["note"] != "":
            print("  note      : " + row["note"][:120])
    if problems == 0:
        print("Nothing - every DOI resolves to the paper you cited.")


